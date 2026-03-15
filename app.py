import streamlit as st
import pandas as pd
from datetime import date, timedelta, datetime
import sqlite3
import hashlib
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import os
import plotly.express as px

# --- 1. إعدادات الصفحة ---
st.set_page_config(page_title="منظومة بيت الشباب - قالمة", layout="wide", page_icon="🏨")

NIGHT_PRICE = 400

# نمط التصميم CSS
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Cairo:wght@400;700&display=swap');
* { font-family: 'Cairo', sans-serif; direction: RTL; text-align: right; }
.main-title { background: linear-gradient(90deg, #1e3c72, #2a5298); color: white; padding: 20px; border-radius: 15px; text-align: center; margin-bottom: 25px; box-shadow: 0 4px 15px rgba(0,0,0,0.2); }
.stat-card { background: white; padding: 20px; border-radius: 15px; border-bottom: 5px solid #1e3c72; text-align: center; box-shadow: 0 4px 12px rgba(0,0,0,0.1); margin-bottom: 10px; }
.red-alert { background-color: #ff4b4b; color: white; padding: 10px; border-radius: 8px; font-weight: bold; text-align: center; margin-bottom: 5px; }
.developer-footer { background: #1e3c72; color: white; padding: 15px; border-radius: 12px; text-align: center; margin-top: 50px; font-size: 0.9rem; }
</style>
""", unsafe_allow_html=True)

# --- 2. إدارة قاعدة البيانات ---
DB_FILE = "hostel_system_v4.db"

def sha256(text): return hashlib.sha256(text.encode("utf-8")).hexdigest()

def get_conn():
    conn = sqlite3.connect(DB_FILE, check_same_thread=False)
    conn.execute("PRAGMA foreign_keys = ON;")
    return conn

def init_db():
    conn = get_conn()
    cur = conn.cursor()
    cur.execute("CREATE TABLE IF NOT EXISTS rooms_config (wing TEXT, room TEXT, beds_count INTEGER, PRIMARY KEY (wing, room))")
    cur.execute("CREATE TABLE IF NOT EXISTS users (role TEXT PRIMARY KEY, password_hash TEXT)")
    cur.execute("""CREATE TABLE IF NOT EXISTS bookings (
        id INTEGER PRIMARY KEY AUTOINCREMENT, full_name TEXT, birth_date DATE, birth_place TEXT,
        address TEXT, id_type TEXT, id_number TEXT, nationality TEXT, visa_date DATE,
        phone_number TEXT, profession TEXT, wing TEXT, room TEXT, bed TEXT,
        minor_doc TEXT, check_in DATE, check_out DATE, payment REAL, status TEXT DEFAULT 'IN', 
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )""")
    cur.execute("""CREATE TABLE IF NOT EXISTS future_bookings (
        id INTEGER PRIMARY KEY AUTOINCREMENT, group_name TEXT, person_count INTEGER, 
        booking_date DATE, phone TEXT
    )""")
    
    if cur.execute("SELECT COUNT(*) FROM users").fetchone()[0] == 0:
        cur.execute("INSERT INTO users VALUES (?,?)", ("مدير", sha256("1234")))
        cur.execute("INSERT INTO users VALUES (?,?)", ("عون استقبال", sha256("5678")))
    
    if cur.execute("SELECT COUNT(*) FROM rooms_config").fetchone()[0] == 0:
        rooms = [("جناح ذكور", f"غرفة {i:02d}", 6) for i in range(1, 6)] + \
                [("جناح إناث", f"غرفة {i:02d}", 6) for i in range(6, 11)]
        cur.executemany("INSERT INTO rooms_config VALUES (?,?,?)", rooms)
    conn.commit()
    conn.close()

init_db()

# --- 3. إدارة الحالة (Session State) ---
if 'booking_data' not in st.session_state: st.session_state.booking_data = {}
if 'step' not in st.session_state: st.session_state.step = "input"
if 'auth' not in st.session_state: st.session_state.auth = False

# --- 4. نظام الدخول ---
if not st.session_state.auth:
    st.markdown('<div class="main-title">🏨 نظام إدارة بيت الشباب محمدي يوسف - قالمة</div>', unsafe_allow_html=True)
    c1, c2, c3 = st.columns([1,2,1])
    with c2:
        role = st.selectbox("اختر الصلاحية", ["مدير", "عون استقبال"])
        pwd = st.text_input("كلمة المرور", type="password")
        if st.button("دخول آمن", use_container_width=True):
            with get_conn() as conn:
                u = conn.execute("SELECT password_hash FROM users WHERE role=?", (role,)).fetchone()
                if u and sha256(pwd) == u[0]:
                    st.session_state.auth, st.session_state.role = True, role
                    st.rerun()
                else: st.error("❌ كلمة المرور خاطئة")
    st.stop()

# --- 5. الواجهة الرئيسية ---
st.sidebar.title(f"👤 {st.session_state.role}")
if st.sidebar.button("🚪 تسجيل الخروج"):
    st.session_state.auth = False
    st.rerun()

tabs = st.tabs(["📊 الإحصائيات", "➕ حجز جديد", "📅 حجوزات مستقبلية", "🔍 السجل والإخلاء", "📁 الأرشيف", "⚙️ الإعدادات"])

# --- التبويب 0: الإحصائيات ---
with tabs[0]:
    with get_conn() as conn:
        total_rooms = conn.execute("SELECT COUNT(*) FROM rooms_config").fetchone()[0]
        # شواغر جناح ذكور
        male_t = conn.execute("SELECT SUM(beds_count) FROM rooms_config WHERE wing='جناح ذكور'").fetchone()[0] or 0
        male_o = conn.execute("SELECT COUNT(*) FROM bookings WHERE status='IN' AND wing='جناح ذكور'").fetchone()[0]
        # شواغر جناح إناث
        female_t = conn.execute("SELECT SUM(beds_count) FROM rooms_config WHERE wing='جناح إناث'").fetchone()[0] or 0
        female_o = conn.execute("SELECT COUNT(*) FROM bookings WHERE status='IN' AND wing='جناح إناث'").fetchone()[0]

    c1, c2, c3 = st.columns(3)
    c1.markdown(f'<div class="stat-card"><h3>{total_rooms}</h3><p>إجمالي عدد الغرف</p></div>', unsafe_allow_html=True)
    c2.markdown(f'<div class="stat-card"><h3>{male_t - male_o}</h3><p>شاغر (جناح ذكور)</p></div>', unsafe_allow_html=True)
    c3.markdown(f'<div class="stat-card"><h3>{female_t - female_o}</h3><p>شاغر (جناح إناث)</p></div>', unsafe_allow_html=True)
    
    # مداخيل
    with get_conn() as conn:
        df_rev = pd.read_sql("SELECT payment, created_at FROM bookings", conn)
    if not df_rev.empty:
        df_rev['month'] = pd.to_datetime(df_rev['created_at']).dt.strftime('%Y-%m')
        st.plotly_chart(px.bar(df_rev.groupby('month')['payment'].sum().reset_index(), x='month', y='payment', title="المداخيل الشهرية"), use_container_width=True)

# --- التبويب 1: حجز جديد ---
with tabs[1]:
    if st.session_state.step == "input":
        st.subheader("📝 إدخال بيانات النزيل")
        with st.container():
            c1, c2 = st.columns(2)
            with c1:
                f_name = st.text_input("الاسم واللقب *", value=st.session_state.booking_data.get('full_name', ''))
                b_date = st.date_input("تاريخ الميلاد *", min_value=date(1900,1,1), max_value=date(2060,12,31), value=st.session_state.booking_data.get('birth_date', date(2000,1,1)))
                b_place = st.text_input("مكان الميلاد *", value=st.session_state.booking_data.get('birth_place', ''))
                address = st.text_input("العنوان الشخصي الإجباري *", value=st.session_state.booking_data.get('address', ''))
                profession = st.text_input("المهنة (اختياري)", value=st.session_state.booking_data.get('profession', ''))
                phone = st.text_input("رقم الهاتف (اختياري)", value=st.session_state.booking_data.get('phone', ''))
            
            with c2:
                id_type = st.selectbox("نوع بطاقة التعريف *", ["بطاقة تعريف عادية", "بيومترية", "رخصة سياقة عادية", "بيومترية", "جواز سفر", "أخرى"])
                id_num = st.text_input("رقم بطاقة التعريف *", value=st.session_state.booking_data.get('id_number', ''))
                nat_type = st.radio("الجنسية *", ["جزائرية", "أجنبي"], horizontal=True)
                
                nat_val = "جزائرية"
                visa_val = None
                if nat_type == "أجنبي":
                    nat_val = st.text_input("الجنسية *")
                    visa_val = st.date_input("تاريخ الدخول (الفيزا) *")
                
                age = (date.today() - b_date).days // 365
                minor_doc = "N/A"
                if age < 18:
                    st.warning(f"⚠️ النزيل قاصر ({age} سنة)")
                    minor_doc = st.selectbox("الوثيقة الإجبارية للقاصر *", ["تصريح أبوي", "حضور الولي", "تصريح من الشرطة", "أخرى"])
                
                with get_conn() as conn:
                    wings = pd.read_sql("SELECT DISTINCT wing FROM rooms_config", conn)['wing'].tolist()
                    w_choice = st.selectbox("الجناح", wings)
                    rooms = pd.read_sql("SELECT room FROM rooms_config WHERE wing=?", conn, params=(w_choice,))['room'].tolist()
                    r_choice = st.selectbox("الغرفة", rooms)
                    beds_max = conn.execute("SELECT beds_count FROM rooms_config WHERE wing=? AND room=?", (w_choice, r_choice)).fetchone()[0]
                    b_choice = st.selectbox("السرير", [f"سرير {i}" for i in range(1, beds_max+1)])
                    nights = st.number_input("عدد الليالي *", min_value=1, value=1)

            if st.button("👁️ مراجعة قبل التأكيد", use_container_width=True):
                if not f_name or not b_place or not address or not id_num or (nat_type=="أجنبي" and not nat_val):
                    st.error("⚠️ يرجى ملء كافة الخانات الإجبارية (*)")
                else:
                    st.session_state.booking_data = {
                        "full_name": f_name, "birth_date": b_date, "birth_place": b_place,
                        "address": address, "id_type": id_type, "id_number": id_num,
                        "nationality": nat_val, "visa_date": visa_val, "phone": phone,
                        "profession": profession, "wing": w_choice, "room": r_choice,
                        "bed": b_choice, "minor_doc": minor_doc, "nights": nights,
                        "payment": nights * NIGHT_PRICE
                    }
                    st.session_state.step = "review"
                    st.rerun()

    elif st.session_state.step == "review":
        st.subheader("📋 مراجعة المعلومات")
        d = st.session_state.booking_data
        col_r1, col_r2 = st.columns(2)
        with col_r1:
            st.info(f"**الاسم:** {d['full_name']}\n\n**الميلاد:** {d['birth_date']} بـ {d['birth_place']}\n\n**العنوان:** {d['address']}")
        with col_r2:
            st.info(f"**الوثيقة:** {d['id_type']} ({d['id_number']})\n\n**الغرفة:** {d['room']} - {d['bed']}\n\n**المبلغ:** {d['payment']} دج")
        
        c_b1, c_b2 = st.columns(2)
        if c_b1.button("🔙 تعديل المعلومات", use_container_width=True):
            st.session_state.step = "input"
            st.rerun()
        if c_b2.button("✅ تأكيد الحجز النهائي", type="primary", use_container_width=True):
            with get_conn() as conn:
                conn.execute("""INSERT INTO bookings 
                    (full_name, birth_date, birth_place, address, id_type, id_number, nationality, visa_date, phone_number, profession, wing, room, bed, minor_doc, check_in, check_out, payment)
                    VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                    (d['full_name'], d['birth_date'], d['birth_place'], d['address'], d['id_type'], d['id_number'], d['nationality'], d['visa_date'], d['phone'], d['profession'], d['wing'], d['room'], d['bed'], d['minor_doc'], date.today(), date.today()+timedelta(days=d['nights']), d['payment']))
                conn.commit()
            st.success("🎉 تم التسجيل بنجاح!"); st.session_state.step = "input"; st.session_state.booking_data = {}; st.rerun()

# --- التبويب 2: حجوزات مستقبلية ---
with tabs[2]:
    st.subheader("📅 الحجوزات القادمة")
    with st.expander("➕ إضافة حجز مستقبلي"):
        with st.form("f_form"):
            fn = st.text_input("اسم الفوج أو الشخص")
            fc = st.number_input("عدد الأشخاص", min_value=1)
            fd = st.date_input("تاريخ الحجز")
            fp = st.text_input("رقم الهاتف")
            if st.form_submit_button("حفظ"):
                with get_conn() as conn:
                    conn.execute("INSERT INTO future_bookings (group_name, person_count, booking_date, phone) VALUES (?,?,?,?)", (fn, fc, fd, fp))
                    conn.commit(); st.rerun()
    
    with get_conn() as conn:
        df_f = pd.read_sql("SELECT * FROM future_bookings ORDER BY booking_date", conn)
    for _, r in df_f.iterrows():
        diff = (r['booking_date'] - date.today()).days
        if diff <= 3: st.markdown(f'<div class="red-alert">🚨 تنبيه: حجز {r["group_name"]} اقترب موعده (بعد {diff} يوم)</div>', unsafe_allow_html=True)
        st.write(f"👤 {r['group_name']} | 👥 {r['person_count']} | 📞 {r['phone']} | 📅 {r['booking_date']}")

# --- التبويب 3: السجل والإخلاء ---
with tabs[3]:
    with get_conn() as conn:
        df_in = pd.read_sql("SELECT id, full_name, room, bed, check_in, check_out, payment FROM bookings WHERE status='IN'", conn)
    st.dataframe(df_in, use_container_width=True)
    sel = st.selectbox("اختر نزيل:", df_in['full_name'] + " (ID: " + df_in['id'].astype(str) + ")", index=None)
    if sel:
        tid = sel.split("(ID: ")[1].replace(")", "")
        c_1, c_2, c_3 = st.columns(3)
        if c_1.button("🚪 إخلاء سبيل"):
            with get_conn() as conn:
                conn.execute("UPDATE bookings SET status='OUT', out_at=? WHERE id=?", (datetime.now(), tid))
                conn.commit(); st.rerun()
        if c_2.button("➕ تمديد ليلة (+400دج)"):
            with get_conn() as conn:
                conn.execute("UPDATE bookings SET check_out = date(check_out, '+1 day'), payment = payment + 400 WHERE id=?", (tid,))
                conn.commit(); st.rerun()
        if c_3.button("🧾 وصل استلام (Word)"):
            doc = Document(); doc.add_heading('وصل إقامة - بيت الشباب قالمة', 0)
            row = df_in[df_in['id']==int(tid)].iloc[0]
            doc.add_paragraph(f"الاسم: {row['full_name']}\nالغرفة: {row['room']}\nتاريخ الدخول: {row['check_in']}\nالمبلغ: {row['payment']} دج")
            b = io.BytesIO(); doc.save(b); st.download_button("⬇️ تحميل الوصل", b.getvalue(), "Receipt.docx")

# --- التبويب 4: الأرشيف ---
with tabs[4]:
    with get_conn() as conn:
        df_arch = pd.read_sql("SELECT full_name, birth_date, birth_place, address, id_type, id_number, phone_number, nights, check_in, created_at, out_at FROM bookings ORDER BY id DESC", conn)
    st.dataframe(df_arch, use_container_width=True)
    
    col_ex1, col_ex2 = st.columns(2)
    with col_ex1:
        out_ex = io.BytesIO(); df_arch.to_excel(out_ex, index=False)
        st.download_button("📊 تصدير للأرشيف Excel", out_ex.getvalue(), "Archive.xlsx")
    with col_ex2:
        if st.button("📝 تصدير الأرشيف لـ Word"):
            doc = Document(); doc.add_heading('سجل النزلاء الكامل', 0)
            table = doc.add_table(rows=1, cols=4); table.style = 'Table Grid'
            for i, h in enumerate(['الاسم واللقب', 'تاريخ ومكان الميلاد', 'العنوان', 'رقم الهوية']): table.rows[0].cells[i].text = h
            for _, r in df_arch.iterrows():
                row = table.add_row().cells
                row[0].text, row[1].text = str(r['full_name']), f"{r['birth_date']} بـ {r['birth_place']}"
                row[2].text, row[3].text = str(r['address']), f"{r['id_type']} ({r['id_number']})"
            b_word = io.BytesIO(); doc.save(b_word); st.download_button("⬇️ تحميل سجل Word", b_word.getvalue(), "Full_Report.docx")

# --- التبويب 5: الإعدادات ---
with tabs[5]:
    st.subheader("⚙️ إعدادات المدير")
    if st.session_state.role == "مدير":
        with st.expander("🔐 تغيير كلمة مرور المدير"):
            new_p = st.text_input("كلمة المرور الجديدة", type="password")
            if st.button("تحديث كلمة السر"):
                with get_conn() as conn:
                    conn.execute("UPDATE users SET password_hash=? WHERE role='مدير'", (sha256(new_p),))
                    conn.commit(); st.success("✅ تم التغيير")
        
        with st.expander("🛠️ إدارة الغرف والأسرة"):
            with get_conn() as conn:
                df_rooms = pd.read_sql("SELECT * FROM rooms_config", conn)
            ed_w = st.selectbox("اختر الجناح", df_rooms['wing'].unique())
            ed_r = st.selectbox("اختر الغرفة", df_rooms[df_rooms['wing']==ed_w]['room'].unique())
            new_bc = st.number_input("عدد الأسرة الجديد", min_value=1, value=6)
            if st.button("تعديل سعة الغرفة"):
                with get_conn() as conn:
                    conn.execute("UPDATE rooms_config SET beds_count=? WHERE wing=? AND room=?", (new_bc, ed_w, ed_r))
                    conn.commit(); st.success("✅ تم التعديل"); st.rerun()
    else: st.warning("هذا التبويب مخصص للمدير فقط")

st.markdown(f'<div class="developer-footer">© RIDHA MERZOUG LABS | <span style="color:#00d4ff;">®ridha_merzoug®</span> - {date.today().year}</div>', unsafe_allow_html=True)
